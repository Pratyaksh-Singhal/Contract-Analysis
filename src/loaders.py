import os
from typing import List
from .interfaces import BaseDocumentLoader, Document


class TextLoader(BaseDocumentLoader):
    def supports(self, file_path: str) -> bool:
        return file_path.lower().endswith(".txt")

    def load(self, file_path: str) -> List[Document]:
        filename = os.path.basename(file_path)
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read().strip()
        if not content:
            return []
        return [Document(content=content, source=filename, page=0, chunk_index=0)]


class PDFLoader(BaseDocumentLoader):
    def supports(self, file_path: str) -> bool:
        return file_path.lower().endswith(".pdf")

    def load(self, file_path: str) -> List[Document]:
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("Install pypdf: pip install pypdf")

        filename = os.path.basename(file_path)
        reader = PdfReader(file_path)
        documents = []
        for page_num, page in enumerate(reader.pages):
            text = (page.extract_text() or "").strip()
            if text:
                documents.append(Document(
                    content=text,
                    source=filename,
                    page=page_num + 1,
                    chunk_index=page_num
                ))
        return documents


class DocxLoader(BaseDocumentLoader):
    def supports(self, file_path: str) -> bool:
        return file_path.lower().endswith(".docx")

    def load(self, file_path: str) -> List[Document]:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")

        filename = os.path.basename(file_path)
        doc = DocxDocument(file_path)
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if not full_text:
            return []
        return [Document(content=full_text, source=filename, page=0, chunk_index=0)]


class LoaderRegistry:
    def __init__(self):
        self._loaders: List[BaseDocumentLoader] = [
            PDFLoader(),
            DocxLoader(),
            TextLoader(),
        ]

    def get_loader(self, file_path: str) -> BaseDocumentLoader:
        for loader in self._loaders:
            if loader.supports(file_path):
                return loader
        raise ValueError(f"No loader found for file: {file_path}")

    def supports(self, file_path: str) -> bool:
        return any(loader.supports(file_path) for loader in self._loaders)
